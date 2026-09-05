"""
Getting text out of an uploaded resume.

Deliberately extract-then-edit: this returns text, and the page puts it in the
textarea for the user to look at before roasting. PDF extraction is lossy and
resume layouts are hostile to it — two columns, tables, icon glyphs — so the
person whose resume it is gets to see what came out. Feeding a mangled
extraction straight to the model produces a roast about the mangling.

Nothing here trusts the filename. The bytes are sniffed, because "resume.pdf"
is a claim made by whoever uploaded the file.
"""

from __future__ import annotations

import io
import re
import zipfile

# Every supported format, as extension -> human label. The dict is also the
# accept list the endpoint and the file picker are built from, so adding a
# format here does not mean remembering three other places.
SUPPORTED = {
    ".pdf": "PDF",
    ".docx": "Word document",
    ".txt": "plain text",
    ".md": "Markdown",
}

MAX_PDF_PAGES = 15

# A resume that extracts to less than this is not a resume we can roast. It is
# almost always a scan: a page of images with no text layer, which pypdf reads
# as empty without complaining.
MIN_USEFUL_CHARS = 120


class ExtractionError(Exception):
    """Could not get usable text out of the file. The message is shown to the user."""


def extract_text(data: bytes, filename: str = "") -> str:
    """
    Text from an uploaded resume, or ExtractionError explaining why not.

    Dispatch is on the bytes rather than the extension. A .docx renamed to
    .pdf is a mistake people genuinely make, and there is no reason to fail
    when the content is perfectly readable.
    """
    if not data:
        raise ExtractionError("That file is empty.")

    if data.startswith(b"%PDF-"):
        return _clean(_from_pdf(data))

    # A .docx is a zip whose first entry is conventionally [Content_Types].xml.
    if data.startswith(b"PK\x03\x04"):
        return _clean(_from_docx(data))

    if _looks_like_text(data):
        return _clean(_from_text(data))

    raise ExtractionError(
        f"That does not look like a {_join(SUPPORTED.values())} file. "
        "Upload one of those, or paste the text instead."
    )


# ------------------------------------------------------------------- formats --


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(io.BytesIO(data))
    except PdfReadError as error:
        raise ExtractionError(f"That PDF could not be read — {error}") from error
    except Exception as error:  # noqa: BLE001
        # pypdf raises whatever the malformed structure happens to produce —
        # ValueError, KeyError, struct.error, a decoder's own exception. None
        # of those should reach the user as a 500 with a traceback.
        raise ExtractionError(
            f"That PDF could not be read ({type(error).__name__}). "
            "It may be damaged — try re-exporting it, or paste the text instead."
        ) from error

    if reader.is_encrypted:
        # An empty password opens the "protected against editing" case, which
        # is what a resume exported from Word with restrictions looks like.
        try:
            opened = reader.decrypt("")
        except Exception:  # noqa: BLE001 - pypdf raises several unrelated types
            opened = 0
        if not opened:
            raise ExtractionError(
                "That PDF is password protected. Remove the password, or paste the text instead."
            )

    pages = reader.pages
    if not len(pages):
        raise ExtractionError("That PDF has no pages in it.")

    # A resume is one or two pages. Anything past the cap is a thesis or a
    # bomb, and either way there is no reason to spend the CPU.
    #
    # Page by page, because one damaged page should not cost the whole
    # document. A resume whose second page has a corrupt font stream is still
    # a resume, and the first page is most of it.
    extracted: list[str] = []
    failed = 0
    for page in pages[:MAX_PDF_PAGES]:
        try:
            extracted.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 - see the constructor note above
            failed += 1

    text = "\n".join(extracted)

    if failed and len(text.strip()) < MIN_USEFUL_CHARS:
        raise ExtractionError(
            "That PDF is damaged — nothing could be read out of it. "
            "Try re-exporting it, or paste the text instead."
        )

    if len(text.strip()) < MIN_USEFUL_CHARS:
        raise ExtractionError(
            "No text came out of that PDF. It is probably a scan or an image — "
            "there is nothing to read. Paste the text instead."
        )
    return text


def _from_docx(data: bytes) -> str:
    import docx
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = docx.Document(io.BytesIO(data))
    except (PackageNotFoundError, KeyError, zipfile.BadZipFile) as error:
        raise ExtractionError(
            "That file is a zip archive but not a Word document. "
            "If it is a .doc, save it as .docx first."
        ) from error
    except Exception as error:  # noqa: BLE001 - same reasoning as the PDF path
        raise ExtractionError(
            f"That Word document could not be read ({type(error).__name__}). "
            "Try re-saving it, or paste the text instead."
        ) from error

    parts = [paragraph.text for paragraph in document.paragraphs]

    # Plenty of resume templates put everything in a borderless table, and
    # document.paragraphs does not reach inside one.
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)

    text = "\n".join(part for part in parts if part.strip())
    if len(text.strip()) < MIN_USEFUL_CHARS:
        raise ExtractionError("That Word document has almost no text in it.")
    return text


def _from_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, UnicodeError):
            continue
    # latin-1 decodes any byte sequence, so reaching here is not possible in
    # practice — but returning something wrong silently would be worse.
    raise ExtractionError("Could not work out the encoding of that text file.")


# ------------------------------------------------------------------ helpers --


# Byte-order marks. UTF-16 has to be recognised before the NUL check below,
# because UTF-16LE encodes every ASCII character as a byte followed by a NUL —
# which is exactly what the binary heuristic is looking for. Windows Notepad
# still writes this when you pick "Unicode" in the save dialog.
_TEXT_BOMS = (b"\xff\xfe", b"\xfe\xff", b"\xef\xbb\xbf")


def _looks_like_text(data: bytes) -> bool:
    """
    Is this plausibly a text file rather than a binary one?

    A NUL byte in the first block is the giveaway for binary — but only once a
    UTF-16 byte-order mark has been ruled out, since UTF-16 is full of them.
    """
    head = data[:4096]
    if head.startswith(_TEXT_BOMS):
        return True
    if b"\x00" in head:
        return False
    try:
        head.decode("utf-8")
        return True
    except UnicodeDecodeError:
        # Could be a truncated multi-byte character at the block boundary, or
        # a legacy encoding. Both are decodable text.
        return sum(byte < 9 for byte in head) == 0


def _clean(text: str) -> str:
    """
    Tidy extractor output without changing what it says.

    PDF extraction leaves a lot of vertical whitespace — one blank line per
    empty layout row — and a resume of mostly blank lines wastes the tokens
    the roast is paid for. Runs of blank lines collapse to one and trailing
    spaces go; nothing else is touched, because reflowing text would change
    what the model sees and therefore what it says.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Non-breaking spaces and the zero-width characters PDF exporters sprinkle
    # in read as ordinary text to a human and as noise to everything else.
    text = text.replace("\xa0", " ").replace("​", "")
    return text.strip()


def _join(labels) -> str:
    labels = list(labels)
    return ", ".join(labels[:-1]) + " or " + labels[-1]
